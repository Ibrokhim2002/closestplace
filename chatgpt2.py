import logging
import io
from aiogram import Bot, Dispatcher, types
from aiogram.contrib.middlewares.logging import LoggingMiddleware
from aiogram.types import ParseMode
from aiogram.utils import executor
from aiogram.dispatcher import FSMContext
import asyncio
import googlemaps
import math
from aiogram.types import KeyboardButton, ReplyKeyboardMarkup, CallbackQuery, User
from telethon.sync import TelegramClient
from telethon.tl.types import InputPeerChannel
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from docx import Document
from docx.shared import Pt
import mysql.connector
import json 
from button import *



# Configure logging
logging.basicConfig(level=logging.INFO)

gmaps = googlemaps.Client(key='AIzaSyBF_1JhujArTeGpdA-Y3MKgyBz3ZnN-Puo')
API_TOKEN = '6160538572:AAF45tZXacT5kwM7lTQs7BLUXiXxMBRk1NQ'


bot = Bot(token='6160538572:AAF45tZXacT5kwM7lTQs7BLUXiXxMBRk1NQ')
storage = MemoryStorage()
dp = Dispatcher(bot, storage=storage)

user_queries = {}  # Dictionary to store user-specific location queries
places_results = {}  # Dictionary to store user-specific place results
x_values = {}  # Dictionary to store user-specific x values


mysql_config = {
    'host': 'us-cdbr-east-06.cleardb.net',
    'user': 'befa09f7c68f53',
    'password': '8d81c04c',
    'database': 'heroku_fdd2ed22f87809d',
    'connect_timeout': 60
}

with open("en.json", "r", encoding="utf-8") as file:
    messages = json.load(file)


chat_ids = set()
location = None
query_type = None
language = "en0"

@dp.message_handler(commands=['start'])
async def start_command(message: types.Message):
    global language
    global messages
    user_id = message.from_user.id
   

    try:
        mysql_connection = mysql.connector.connect(**mysql_config)
        mysql_cursor = mysql_connection.cursor()

    except mysql.connector.Error as error:
        print('Error connecting to the database:', error)

    # Check if the user ID exists in the database
    select_query = "SELECT language FROM users_tg WHERE user_id = %s"
    select_values = (user_id,)
    mysql_cursor.execute(select_query, select_values)

    existing_user = mysql_cursor.fetchone()

    if existing_user is None:
        # User ID does not exist, prompt for language selection
        await bot.send_message(chat_id=message.chat.id, text="Please choose your language:",
                               reply_markup=buttons_column2)
    else:
        # User ID already exists, retrieve the language and welcome message
        language = existing_user[0]



        # Load the language-specific messages from the JSON file
        with open("en.json", "r", encoding="utf-8") as file:
            messages = json.load(file)

        user_location, buttons_column1, more1 = language_user(language, messages)
        await bot.send_message(chat_id=message.chat.id, text=messages[int(existing_user[-1][-1])][language]["welcome_message"], reply_markup=buttons_column1)

    mysql_cursor.close()
    mysql_connection.close()
    language_user(language, messages)





@dp.callback_query_handler(lambda callback_query: callback_query.data in ['uz1', 'ru2', 'en0'])
async def language_selected(callback_query: types.CallbackQuery):
    global user_queries
    global places_results
    global location
    global query_type
    global language
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    user_id = callback_query.from_user.id
    username = callback_query.from_user.username
    first_name = callback_query.from_user.first_name
    last_name = callback_query.from_user.last_name

    try:
        mysql_connection = mysql.connector.connect(**mysql_config)
        mysql_cursor = mysql_connection.cursor()

    except mysql.connector.Error as error:
        print('Error connecting to the database:', error)

    # Check if the user ID already exists in the database
    select_query = "SELECT user_id FROM users_tg WHERE user_id = %s"
    select_values = (user_id,)
    mysql_cursor.execute(select_query, select_values)

    existing_user = mysql_cursor.fetchone()

    if existing_user is None:
        # User ID does not exist, insert a new record
        insert_query = "INSERT INTO users_tg (user_id, username, first_name, last_name, language) VALUES (%s, %s, %s, %s, %s)"
        insert_values = (user_id, username, first_name, last_name, callback_query.data)
        mysql_cursor.execute(insert_query, insert_values)
        mysql_connection.commit()
    else:
        # User ID already exists
        update_query = "UPDATE users_tg SET language = %s WHERE user_id = %s"
        update_values = (callback_query.data, user_id)
        mysql_cursor.execute(update_query, update_values)
        mysql_connection.commit()

    chat_ids.add(callback_query.message.chat.id)

    select_query = "SELECT language FROM users_tg WHERE user_id = %s"
    select_values = (user_id,)
    mysql_cursor.execute(select_query, select_values)

    # Fetch the result
    result = mysql_cursor.fetchone()


    # Retrieve the selected language
    language = callback_query.data

    # Load the language-specific messages from the JSON file
    with open("en.json", "r", encoding="utf-8") as file:
        messages = json.load(file)

    user_location, buttons_column1, more1 = language_user(language, messages)
    await bot.send_message(chat_id=user_id, text=messages[int(result[-1][-1])][language]["welcome_message"], reply_markup=buttons_column1)

    mysql_cursor.close()
    mysql_connection.close()






class User:
    def __init__(self, id, username, first_name, last_name):
        self.id = id
        self.username = username
        self.first_name = first_name
        self.last_name = last_name



async def export_users():
    mysql_connection = mysql.connector.connect(**mysql_config)
    # Create a MySQL cursor
    mysql_cursor = mysql_connection.cursor()
    num = 0

    # Retrieve user data from the database
    select_data_query = "SELECT * FROM users_tg"
    mysql_cursor.execute(select_data_query)
    results = mysql_cursor.fetchall()

    # Count the total number of users
    total_users_query = "SELECT COUNT(*) FROM users_tg"
    mysql_cursor.execute(total_users_query)
    total_users = mysql_cursor.fetchone()[0]

    # Create a new Word document
    doc = Document()
    doc.add_heading('User Data', level=1)

    # Create a table and add headers
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'User ID'
    header_cells[2].text = 'Username'
    header_cells[3].text = 'First Name'
    header_cells[4].text = 'Last Name'
    for cell in header_cells:
        cell.paragraphs[0].paragraph_format.alignment = 1  # Center alignment
        
    # Add user data to the table
    for row in results:
        num += 1
        row_cells = table.add_row().cells
        row_cells[0].text = str(num)
        row_cells[1].text = str(row[0])
        row_cells[2].text = row[1] if row[1] else ''
        row_cells[3].text = row[2] if row[2] else ''
        row_cells[4].text = row[3] if row[3] else ''
        for cell in row_cells:
            cell.width = Pt(100)
            cell.paragraphs[0].paragraph_format.alignment = 1  # Center alignment

    # Add total number of users at the bottom
    total_row = table.add_row().cells
    total_row[0].merge(total_row[4])
    total_row[0].text = f"Total Users: {total_users}"
    total_row[0].paragraphs[0].paragraph_format.alignment = 1  # Center alignment

    # Adjust cell padding and vertical alignment
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.vertical_alignment = 1  # Center alignment

    # Save the document
    doc.save('user_data.docx')
    mysql_cursor.close()
    mysql_connection.close()




@dp.message_handler(commands=['exportusers'])
async def handle_export_users(message: types.Message):
    # Check if the user is an admin
    if message.from_user.id != 398222109:
        return

    # Export and send the users document
    await export_users()
    with open('user_data.docx', 'rb') as file:
        await bot.send_document(chat_id=message.chat.id, document=file)



@dp.message_handler(commands=['lng'])
async def start_command(message: types.Message):
    await bot.send_message(chat_id=message.chat.id, text="Please choose your language:",
                               reply_markup=buttons_column2)




@dp.message_handler()
async def send_welcome2(message: types.Message):
    global user_queries
    global messages
    global language
    user_queries[message.chat.id] = message.text
    user_location, buttons_column1, more1 = language_user(language, messages)   
    await bot.send_message(chat_id=message.chat.id, text=messages[int(language[-1][-1])][language]["location_query_prompt"], reply_markup=user_location)



@dp.message_handler(content_types=types.ContentType.LOCATION)
async def handle_message(message: types.Message):
    global user_queries
    global places_results
    global x_values
    global location
    global query_type
    global language
    global messages

    location = message.location
    lat, lng = location.latitude, location.longitude
    chat_id = message.chat.id


    # Check if the chat_id exists in user_queries only if query_type is None
    if query_type is None and chat_id not in user_queries:
        await bot.send_message(chat_id=chat_id, text=messages[int(language[-1][-1])][language]["query_before"])
        return

    query = user_queries.get(chat_id, "")

    if query_type is not None:
        places_result = gmaps.places_nearby(location=(lat, lng), radius=5000, type=query_type)
        # Reset the query_type value after using it
        query_type = None
    else:
        places_result = gmaps.places_nearby(location=(lat, lng), radius=5000, keyword=query)
    #places_result = gmaps.places_nearby(location=(lat, lng), radius=5000, keyword=query)
    #places_result = gmaps.places_nearby(**search_params)


    if len(places_result['results']) == 0:
        await bot.send_message(chat_id=chat_id, text=messages[int(language[-1][-1])][language]["no_places_found"])
        return

    closest_places = places_result['results']
    closest_distance = float('inf')
    items = []

    for place in places_result['results']:
        mosque_lat = place['geometry']['location']['lat']
        mosque_lng = place['geometry']['location']['lng']
        distance_meters = calculate_distance(lat, lng, mosque_lat, mosque_lng)

        if distance_meters < closest_distance:
            closest_distance = distance_meters
        items.append((place, distance_meters))

    # Sort the places based on distance in ascending order
    sorted_places = sorted(items, key=lambda x: x[1])

    places_results[chat_id] = sorted_places
    x_values[chat_id] = 1 # Initialize x to 0 for new location search

    if sorted_places:
        for place, distance in sorted_places:
            photo_id = place.get('photos', [{}])[0].get('photo_reference', '')
            if photo_id:
                raw_image_data = gmaps.places_photo(photo_reference=photo_id, max_height=400, max_width=400)


                caption = messages[int(language[-1][-1])][language]["place_info"].format(name=place['name'], vicinity=place['vicinity'], distance=distance)
                photo_bytes = b''.join(chunk for chunk in raw_image_data if chunk)

                place_lat = place['geometry']['location']['lat']
                place_lng = place['geometry']['location']['lng']
                user_location, buttons_column1, more1 = language_user(language, messages)
                mosque_location = await bot.send_venue(
                    chat_id=message.chat.id,
                    latitude=place_lat,
                    longitude=place_lng,
                    title=place['name'],
                    address=place['vicinity'],
                    reply_markup=more1
                )
                await bot.send_photo(chat_id=message.chat.id, photo=photo_bytes, caption=caption, parse_mode=ParseMode.HTML)

                break  # Only send the first place for now, additional places can be retrieved using the "More" button
            else:
                # No place with a photo found, send the location of the closest place
                place, distance = sorted_places[0]
                place_lat = place['geometry']['location']['lat']
                place_lng = place['geometry']['location']['lng']
                user_location, buttons_column1, more1 = language_user(language, messages)
                mosque_location = await bot.send_venue(
                    chat_id=message.chat.id,
                    latitude=place_lat,
                    longitude=place_lng,
                    title=place['name'],
                    address=place['vicinity'],
                    reply_markup=more1
                )
                await bot.send_message(message.chat.id, messages[int(language[-1][-1])][language]["no_photo_found"].format(name=place['name'], vicinity=place['vicinity'], distance=distance))
                break
    else:
        # No places found at all
        await bot.send_message(chat_id=chat_id, text=messages[int(language[-1][-1])][language]["no_places_found"])




# Define the callback function for the "More" button
async def button_callback(call: types.CallbackQuery):
    global places_results
    global x_values
    global messages
    global language

    chat_id = call.message.chat.id

    if chat_id not in places_results or chat_id not in x_values:
        await bot.send_message(chat_id=chat_id, text=messages[int(language[-1][-1])][language]["no_more_places_available"])
        return

    sorted_places = places_results[chat_id]
    x = x_values[chat_id]
    if x < len(sorted_places):
        x += 1  # Increment x to start from index 1
        place, distance = sorted_places[x - 1]  # Subtract 1 to access correct index

        photo_id = place.get('photos', [{}])[0].get('photo_reference', '')
        if photo_id:
            raw_image_data = gmaps.places_photo(photo_reference=photo_id, max_height=400, max_width=400)
            caption = messages[int(language[-1][-1])][language]["place_info"].format(name=place['name'], vicinity=place['vicinity'], distance=distance)
            photo_bytes = b''.join(chunk for chunk in raw_image_data if chunk)
            place_lat = place['geometry']['location']['lat']
            place_lng = place['geometry']['location']['lng']
            user_location, buttons_column1, more1 = language_user(language, messages)
            mosque_location = await bot.send_venue(
                chat_id=call.message.chat.id,
                latitude=place_lat,
                longitude=place_lng,
                title=place['name'],
                address=place['vicinity'],
                reply_markup=more1
            )
            await bot.send_photo(chat_id=call.message.chat.id, photo=photo_bytes, caption=caption, parse_mode=types.ParseMode.HTML)
            
        else:
            caption = messages[int(language[-1][-1])][language]["place_info"].format(name=place['name'], vicinity=place['vicinity'], distance=distance)
            place_lat = place['geometry']['location']['lat']
            place_lng = place['geometry']['location']['lng']
            user_location, buttons_column1, more1 = language_user(language, messages)
            mosque_location = await bot.send_venue(
                chat_id=call.message.chat.id,
                latitude=place_lat,
                longitude=place_lng,
                title=place['name'],
                address=place['vicinity'],
                reply_markup=more1
            )
            await bot.send_message(call.message.chat.id, messages[int(language[-1][-1])][language]["no_photo_found"].format(name=place['name'], vicinity=place['vicinity'], distance=distance))
            

        x_values[chat_id] = x
    else:
        await bot.send_message(call.message.chat.id, messages[int(language[-1][-1])][language]["no_more_places_available"])





# Example usage of the "More" button and callback
@dp.callback_query_handler(lambda call: call.data == "more")
async def btn(call: types.CallbackQuery):
    await button_callback(call)




def calculate_distance(lat1, lon1, lat2, lon2):
    R = 6371  # Radius of the Earth in kilometers

    dlat = math.radians(lat2 - lat1)
    dlon = math.radians(lon2 - lon1)

    a = math.sin(dlat / 2) * math.sin(dlat / 2) + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlon / 2) * math.sin(dlon / 2)
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))

    distance_km = R * c
    distance_meters = distance_km * 1000

    return distance_meters






############################################################   BUTTONS  #################################################################

@dp.callback_query_handler(text="airport")
async def btn(call: CallbackQuery):
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global location
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "airport"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return

@dp.callback_query_handler(text="bank")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "bank"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return

@dp.callback_query_handler(text="mosque")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "mosque"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return


@dp.callback_query_handler(text="gas_station")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "gas_station"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return

@dp.callback_query_handler(text="amusement_park")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "amusement_park"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return


@dp.callback_query_handler(text="supermarket")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "supermarket"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return

@dp.callback_query_handler(text="hospital")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "hospital"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return



@dp.callback_query_handler(text="restaurant")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "restaurant"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return



@dp.callback_query_handler(text="subway_station")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "subway_station"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return


@dp.callback_query_handler(text="pharmacy")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "pharmacy"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return


@dp.callback_query_handler(text="car_wash")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "car_wash"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return


@dp.callback_query_handler(text="bus_station")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    if location is None:
        query_type = "bus_station"
        user_location, buttons_column1, more1 = language_user(language, messages)
        await call.message.answer(messages[int(language[-1][-1])][language]["share_location"], reply_markup=user_location)
        return


@dp.callback_query_handler(text="type_yourself")
async def btn(call: CallbackQuery):
    global location
    global location
    global user_queries
    global places_results
    global location
    global query_type
    global messages

    user_queries = {}  # Reset user_queries dictionary
    places_results = {}  # Reset places_results dictionary
    location = None  # Reset location variable
    query_type = None  # Reset query_type variable

    await call.message.answer(messages[int(language[-1][-1])][language]["place_name_prompt"])


if __name__ == '__main__':
    executor.start_polling(dp, skip_updates=True)






